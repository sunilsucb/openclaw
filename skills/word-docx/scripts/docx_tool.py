#!/usr/bin/env python3
# /// script
# requires-python = ">=3.10"
# dependencies = [
#     "python-docx>=1.0.0",
# ]
# ///
"""
DOCX utility: create, read, inspect, and validate Word documents.

Usage:
    uv run docx_tool.py read <file.docx>                     # Extract text
    uv run docx_tool.py inspect <file.docx>                  # Show structure (styles, sections, etc.)
    uv run docx_tool.py create <file.docx> --title "Title"   # Create new document
    uv run docx_tool.py validate <file.docx>                 # Safety checks (zip bomb, macros, etc.)
"""

import argparse
import json
import os
import sys
import zipfile
from pathlib import Path

# Safety limits
MAX_FILE_SIZE = 100 * 1024 * 1024       # 100MB max input file
MAX_ZIP_RATIO = 100                      # Max decompression ratio (zip bomb detection)
MAX_ZIP_ENTRIES = 5000                   # Max entries in the ZIP archive
DANGEROUS_EXTENSIONS = {".docm", ".dotm", ".doc"}


def die(msg: str) -> None:
    print(f"Error: {msg}", file=sys.stderr)
    sys.exit(1)


def validate_input_path(path_str: str) -> Path:
    """Validate input file exists and is within size limits."""
    p = Path(path_str)
    if not p.exists():
        die(f"File not found: {path_str}")
    if not p.is_file():
        die(f"Not a file: {path_str}")
    size = p.stat().st_size
    if size > MAX_FILE_SIZE:
        die(f"File too large ({size:,} bytes, max {MAX_FILE_SIZE:,}).")
    if size == 0:
        die("File is empty (0 bytes).")
    return p


def validate_output_path(path_str: str) -> Path:
    """Validate output path doesn't traverse directories unsafely."""
    p = Path(path_str)
    resolved = p.resolve()
    cwd = Path.cwd().resolve()
    try:
        resolved.relative_to(cwd)
    except ValueError:
        die(f"Output path must be within current directory ({cwd}).")
    return p


def check_zip_safety(filepath: Path) -> list[str]:
    """Check for zip bombs and suspicious content."""
    warnings = []

    try:
        with zipfile.ZipFile(filepath, 'r') as zf:
            entries = zf.infolist()

            if len(entries) > MAX_ZIP_ENTRIES:
                warnings.append(f"CRITICAL: Too many ZIP entries ({len(entries)}, max {MAX_ZIP_ENTRIES}) — possible zip bomb.")
                return warnings

            compressed_total = sum(e.compress_size for e in entries)
            uncompressed_total = sum(e.file_size for e in entries)

            if compressed_total > 0:
                ratio = uncompressed_total / compressed_total
                if ratio > MAX_ZIP_RATIO:
                    warnings.append(f"CRITICAL: High decompression ratio ({ratio:.1f}x, max {MAX_ZIP_RATIO}x) — possible zip bomb.")

            # Check for macro-bearing content
            for entry in entries:
                name_lower = entry.filename.lower()
                if "vbaproject" in name_lower or name_lower.endswith(".bin"):
                    warnings.append(f"WARNING: Macro content detected ({entry.filename}). File may contain executable code.")
                if entry.filename.startswith("/") or ".." in entry.filename:
                    warnings.append(f"CRITICAL: Path traversal in ZIP entry ({entry.filename}).")

    except zipfile.BadZipFile:
        warnings.append("CRITICAL: Not a valid ZIP file (corrupt or not a real .docx).")
    except Exception as e:
        warnings.append(f"ERROR: Could not inspect ZIP: {e}")

    return warnings


def cmd_validate(args):
    """Validate a .docx file for safety issues."""
    filepath = validate_input_path(args.file)
    ext = filepath.suffix.lower()

    results = {
        "file": str(filepath),
        "size_bytes": filepath.stat().st_size,
        "extension": ext,
        "warnings": [],
        "safe": True,
    }

    if ext in DANGEROUS_EXTENSIONS:
        results["warnings"].append(f"WARNING: {ext} files may contain macros or use legacy format.")

    zip_warnings = check_zip_safety(filepath)
    results["warnings"].extend(zip_warnings)

    if any(w.startswith("CRITICAL") for w in results["warnings"]):
        results["safe"] = False

    print(json.dumps(results, indent=2))
    if not results["safe"]:
        sys.exit(1)


def cmd_read(args):
    """Extract text content from a .docx file."""
    filepath = validate_input_path(args.file)

    # Safety check first
    zip_warnings = check_zip_safety(filepath)
    critical = [w for w in zip_warnings if w.startswith("CRITICAL")]
    if critical:
        for w in critical:
            print(w, file=sys.stderr)
        die("File failed safety checks. Run 'validate' for details.")

    from docx import Document

    try:
        doc = Document(str(filepath))
    except Exception as e:
        die(f"Failed to open document: {e}")

    output = {
        "file": str(filepath),
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "sections": len(doc.sections),
        "text": [],
    }

    for para in doc.paragraphs:
        entry = {"text": para.text}
        if para.style and para.style.name:
            entry["style"] = para.style.name
        output["text"].append(entry)

    # Extract table content
    if doc.tables:
        output["table_data"] = []
        for i, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                cells = [cell.text for cell in row.cells]
                rows.append(cells)
            output["table_data"].append({"table_index": i, "rows": rows})

    print(json.dumps(output, indent=2, ensure_ascii=False))


def cmd_inspect(args):
    """Inspect document structure."""
    filepath = validate_input_path(args.file)

    zip_warnings = check_zip_safety(filepath)
    critical = [w for w in zip_warnings if w.startswith("CRITICAL")]
    if critical:
        for w in critical:
            print(w, file=sys.stderr)
        die("File failed safety checks.")

    from docx import Document

    try:
        doc = Document(str(filepath))
    except Exception as e:
        die(f"Failed to open document: {e}")

    # Collect styles in use
    styles_used = set()
    for para in doc.paragraphs:
        if para.style and para.style.name:
            styles_used.add(para.style.name)

    # Section info
    sections = []
    for i, section in enumerate(doc.sections):
        sections.append({
            "index": i,
            "page_width_inches": round(section.page_width.inches, 2) if section.page_width else None,
            "page_height_inches": round(section.page_height.inches, 2) if section.page_height else None,
            "left_margin_inches": round(section.left_margin.inches, 2) if section.left_margin else None,
            "right_margin_inches": round(section.right_margin.inches, 2) if section.right_margin else None,
            "orientation": str(section.orientation) if section.orientation else None,
        })

    # ZIP contents
    zip_parts = []
    try:
        with zipfile.ZipFile(filepath, 'r') as zf:
            for entry in zf.infolist():
                zip_parts.append({
                    "name": entry.filename,
                    "size": entry.file_size,
                })
    except Exception:
        pass

    output = {
        "file": str(filepath),
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "sections": sections,
        "styles_used": sorted(styles_used),
        "zip_parts": zip_parts,
        "warnings": zip_warnings,
    }

    print(json.dumps(output, indent=2, ensure_ascii=False))


def cmd_create(args):
    """Create a new .docx document."""
    output_path = validate_output_path(args.file)

    from docx import Document
    from docx.shared import Pt

    doc = Document()

    if args.title:
        doc.add_heading(args.title, level=0)

    if args.body:
        for line in args.body:
            doc.add_paragraph(line)
    else:
        doc.add_paragraph("")  # Empty starter paragraph

    # Set default font if specified
    if args.font:
        style = doc.styles['Normal']
        font = style.font
        font.name = args.font

    if args.font_size:
        style = doc.styles['Normal']
        font = style.font
        font.size = Pt(args.font_size)

    output_path.parent.mkdir(parents=True, exist_ok=True)

    try:
        doc.save(str(output_path))
    except Exception as e:
        die(f"Failed to save document: {e}")

    print(json.dumps({
        "created": str(output_path.resolve()),
        "title": args.title,
        "paragraphs": len(doc.paragraphs),
    }, indent=2))


def main():
    parser = argparse.ArgumentParser(description="DOCX utility tool")
    subparsers = parser.add_subparsers(dest="command", required=True)

    # read
    p_read = subparsers.add_parser("read", help="Extract text from a .docx file")
    p_read.add_argument("file", help="Path to .docx file")

    # inspect
    p_inspect = subparsers.add_parser("inspect", help="Inspect document structure")
    p_inspect.add_argument("file", help="Path to .docx file")

    # validate
    p_validate = subparsers.add_parser("validate", help="Safety-check a .docx file")
    p_validate.add_argument("file", help="Path to .docx file")

    # create
    p_create = subparsers.add_parser("create", help="Create a new .docx document")
    p_create.add_argument("file", help="Output .docx path")
    p_create.add_argument("--title", "-t", help="Document title")
    p_create.add_argument("--body", "-b", nargs="*", help="Body paragraphs")
    p_create.add_argument("--font", help="Default font name (e.g., 'Calibri')")
    p_create.add_argument("--font-size", type=float, help="Default font size in points")

    args = parser.parse_args()

    commands = {
        "read": cmd_read,
        "inspect": cmd_inspect,
        "validate": cmd_validate,
        "create": cmd_create,
    }
    commands[args.command](args)


if __name__ == "__main__":
    main()
